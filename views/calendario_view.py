import flet as ft
from controllers.transmissao_controller import TransmissaoController
from datetime import datetime, timedelta
from utils.helpers import normalize_date, format_date_br, formatar_data_completa_semana, parse_date, formatar_data_semana, get_status_info, get_status_options
from utils.dialog_helper import mostrar_detalhes_transmissao, abrir_menu_contexto
import asyncio
import os
import tkinter as tk
from tkinter import filedialog
import calendar

try:
    from docx import Document; from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from openpyxl import Workbook; from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
except ImportError: pass

class CalendarioView(ft.Column):
    def __init__(self, controller: TransmissaoController, on_edit):
        super().__init__(expand=True, spacing=25)
        self.controller = controller
        self.on_edit = on_edit
        self.pagina_atual = 0
        self.itens_por_página = 20
        self.view_mode = "calendario"
        self.selection_mode = False
        self.selected_ids = set()
        
        
        self.txt_busca = ft.TextField(
            hint_text="Pesquisar...", 
            prefix_icon=ft.Icons.SEARCH, 
            on_change=lambda _: self.reset_and_update(), 
            width=220, height=45, text_size=13, 
            content_padding=ft.padding.only(left=10, right=10), 
            border_radius=10, 
            bgcolor=ft.Colors.with_opacity(0.1, ft.Colors.WHITE),
            border_color=ft.Colors.WHITE10,
            focused_border_color=ft.Colors.CYAN_400
        )
        self.dd_periodo = ft.Dropdown(
            label="Período", 
            options=[ft.dropdown.Option("Mês"), ft.dropdown.Option("Personalizado")], 
            value="Mês", 
            on_select=self.on_periodo_change, 
            width=160
        )
        
        self.meses_nomes = ["Janeiro", "Fevereiro", "Março", "Abril", "Maio", "Junho", "Julho", "Agosto", "Setembro", "Outubro", "Novembro", "Dezembro"]
        self.row_meses_disponiveis = ft.Row(spacing=15)
        
        self.txt_mes_selecionado = ft.Text(self.meses_nomes[datetime.now().month-1], color=ft.Colors.WHITE, size=18, weight=ft.FontWeight.BOLD)
        self.menu_mes = ft.PopupMenuButton(
            content=ft.Row([self.txt_mes_selecionado, ft.Icon(ft.Icons.ARROW_DROP_DOWN, color=ft.Colors.WHITE, size=20)], spacing=5),
            items=[ft.PopupMenuItem(content=ft.Text("Todos"), on_click=lambda e: self.set_filtro_mes_bar(0, "Todos"))] + \
                  [ft.PopupMenuItem(content=ft.Text(m), on_click=lambda e, m=m, i=i: self.set_filtro_mes_bar(i+1, m)) for i, m in enumerate(self.meses_nomes)]
        )
        
        self.txt_ano_selecionado = ft.Text(str(datetime.now().year), color=ft.Colors.WHITE, size=18, weight=ft.FontWeight.BOLD)
        self.menu_ano = ft.PopupMenuButton(
            content=ft.Row([self.txt_ano_selecionado, ft.Icon(ft.Icons.ARROW_DROP_DOWN, color=ft.Colors.WHITE, size=20)], spacing=5),
            items=[]
        )

        self.filter_bar = ft.Container(
            content=ft.Row([
                ft.Container(content=ft.Row([ft.Text("MESES:", color=ft.Colors.WHITE60, size=11, weight=ft.FontWeight.BOLD), self.row_meses_disponiveis], spacing=15), expand=True),
                ft.Container(content=self.menu_mes, alignment=ft.Alignment(0, 0)),
                ft.Container(
                    content=self.menu_ano, 
                    alignment=ft.Alignment(1, 0),
                    border=ft.border.all(1, ft.Colors.WHITE38),
                    border_radius=8,
                    padding=ft.padding.symmetric(horizontal=12, vertical=4),
                    bgcolor=ft.Colors.with_opacity(0.1, ft.Colors.BLACK)
                )
            ], alignment=ft.MainAxisAlignment.SPACE_BETWEEN),
            gradient=ft.LinearGradient(
                begin=ft.Alignment(-1, 0),
                end=ft.Alignment(1, 0),
                colors=[ft.Colors.GREEN_900, ft.Colors.GREEN_600]
            ),
            padding=ft.padding.symmetric(horizontal=20, vertical=8),
            border_radius=10,
            margin=ft.margin.only(bottom=5),
            visible=True
        )
        
        self.dd_filtro_mes = ft.Dropdown(
            label="Mês",
            options=[ft.dropdown.Option(text=m, key=str(i+1)) for i, m in enumerate(self.meses_nomes)],
            value=str(datetime.now().month),
            on_select=lambda _: self.reset_and_update(),
            width=130, height=45, text_size=12, visible=False
        )
        
        anos_disp = sorted(list(set(parse_date(t.data).year for t in self.controller.transmissoes if parse_date(t.data))), reverse=True)
        if not anos_disp: anos_disp = [datetime.now().year]
        
        self.dd_filtro_ano = ft.Dropdown(
            label="Ano",
            options=[ft.dropdown.Option(str(a)) for a in anos_disp],
            value=str(datetime.now().year),
            on_select=lambda _: self.reset_and_update(),
            width=100, height=45, text_size=12, visible=False
        )
        
        self.txt_data_inicio = ft.TextField(
            label="De:", hint_text="dd/mm/aaaa", 
            width=120, height=45, text_size=12, 
            on_submit=lambda _: self.reset_and_update(),
            visible=False,
            bgcolor=ft.Colors.with_opacity(0.15, ft.Colors.WHITE),
            border_radius=8,
            border_color=ft.Colors.WHITE10
        )
        self.txt_data_fim = ft.TextField(
            label="Até:", hint_text="dd/mm/aaaa", 
            width=120, height=45, text_size=12, 
            on_submit=lambda _: self.reset_and_update(),
            visible=False,
            bgcolor=ft.Colors.with_opacity(0.15, ft.Colors.WHITE),
            border_radius=8,
            border_color=ft.Colors.WHITE10
        )

        self.dd_status = ft.Dropdown(label="Filtrar por Status", options=[ft.dropdown.Option("Todos")] + [ft.dropdown.Option(o) for o in get_status_options()], value="Todos", width=170, on_select=lambda _: self.reset_and_update())
        tipos = sorted(list(set(t.tipo_transmissao for t in self.controller.transmissoes if t.tipo_transmissao)))
        self.dd_tipo = ft.Dropdown(label="Tipo", options=[ft.dropdown.Option("Todos")] + [ft.dropdown.Option(o) for o in tipos], value="Todos", width=120, on_select=lambda _: self.reset_and_update())
        
        modalidades = sorted(list(set(t.modalidade for t in self.controller.transmissoes if t.modalidade)))
        self.dd_modalidade = ft.Dropdown(label="Modalidade", options=[ft.dropdown.Option("Todos")] + [ft.dropdown.Option(o) for o in modalidades], value="Todos", width=140, on_select=lambda _: self.reset_and_update())
        self.dd_ordem = ft.Dropdown(label="Ordenar por", options=[ft.dropdown.Option("Data do Evento"), ft.dropdown.Option("Nome do Evento")], value="Data do Evento", width=200, on_select=lambda _: self.reset_and_update())
        
        self.btn_aplicar = ft.ElevatedButton(
            "Aplicar", 
            icon=ft.Icons.FILTER_ALT, 
            on_click=lambda _: self.reset_and_update(), 
            style=ft.ButtonStyle(bgcolor=ft.Colors.GREEN_700, color=ft.Colors.WHITE),
            height=45
        )
        
        self.dd_view_mode = ft.Dropdown(
            label="Modo",
            options=[
                ft.dropdown.Option("calendario", "📅 Calendário"),
                ft.dropdown.Option("list", "📋 Lista"),
                ft.dropdown.Option("grid", "🃏 Cards"),
            ],
            value=self.view_mode,
            on_select=self.on_view_mode_change,
            width=160, height=45, text_size=13,
            color=ft.Colors.CYAN_400,
            border_color=ft.Colors.CYAN_400,
            focused_border_color=ft.Colors.CYAN_200,
            label_style=ft.TextStyle(color=ft.Colors.CYAN_200, weight=ft.FontWeight.BOLD),
        )
        self.txt_count = ft.Text("0 eventos encontrados", color=ft.Colors.WHITE70, size=12)
        self.grid_container = ft.Column(expand=True, scroll=ft.ScrollMode.ADAPTIVE)
        
        self.btn_exportar_excel = ft.ElevatedButton("Excel", icon=ft.Icons.FILE_DOWNLOAD, on_click=self.exportar_excel, style=ft.ButtonStyle(color=ft.Colors.WHITE, bgcolor=ft.Colors.GREEN_700), height=40)
        self.btn_exportar_links = ft.ElevatedButton("Links", icon=ft.Icons.LINK, on_click=self.exportar_links, style=ft.ButtonStyle(color=ft.Colors.WHITE, bgcolor=ft.Colors.BLUE_700), height=40)
        self.btn_cancelar_links = ft.TextButton("Cancelar", icon=ft.Icons.CLOSE, on_click=self.cancelar_selecao_links, visible=False, style=ft.ButtonStyle(color=ft.Colors.RED_400))

        self.controls = [
            ft.Stack([
                ft.Column([ft.Text("Calendário de Transmissões", size=28, weight=ft.FontWeight.BOLD), self.txt_count], spacing=2),
                ft.Container(
                    content=ft.Row([self.btn_cancelar_links, self.btn_exportar_excel, self.btn_exportar_links], spacing=10, tight=True),
                    alignment=ft.Alignment(1, -1)
                )
            ]),
            ft.Container(height=20),
            ft.Row([
                self.txt_busca, 
                self.dd_periodo,
                self.dd_filtro_mes,
                self.dd_filtro_ano,
                self.txt_data_inicio,
                self.txt_data_fim,
                self.dd_status, 
                self.dd_tipo, 
                self.dd_modalidade, 
                self.dd_ordem, 
                self.dd_view_mode
            ], spacing=10, wrap=True),

            self.filter_bar,
            ft.Divider(height=1, color=ft.Colors.WHITE10),
            self.grid_container
        ]
        self.update_year_menu()
        self.on_periodo_change(None)
        # self.init_ui() # Removido para evitar chamada duplicada (já chamado em on_periodo_change)

    def on_view_mode_change(self, e):
        self.view_mode = self.dd_view_mode.value
        self.init_ui()

    def on_periodo_change(self, e):
        p = self.dd_periodo.value
        self.dd_filtro_mes.visible = False
        self.dd_filtro_ano.visible = False
        self.txt_data_inicio.visible = (p == "Personalizado")
        self.txt_data_fim.visible = (p == "Personalizado")
        self.filter_bar.visible = (p in ["Mês", "Ano"])
        self.reset_and_update()

    def get_lista_filtrada(self):
        hoje = datetime.now()
        hoje_str = hoje.strftime("%Y-%m-%d")
        termo = self.txt_busca.value.lower() if self.txt_busca.value else ""
        periodo = self.dd_periodo.value
        
        # Valores dos sub-filtros
        mes_f = int(self.dd_filtro_mes.value) if self.dd_filtro_mes.value else hoje.month
        ano_f = int(self.dd_filtro_ano.value) if self.dd_filtro_ano.value else hoje.year
        dt_ini = None; dt_fim = None

        # Condição para ocultar finalizados no estado inicial/padrão
        is_default_view = (
            periodo == "Mês" and 
            mes_f == hoje.month and 
            ano_f == hoje.year and 
            not termo and 
            self.dd_status.value == "Todos" and 
            self.dd_tipo.value == "Todos" and 
            self.dd_modalidade.value == "Todos"
        )
        if periodo == "Personalizado":
            try:
                if self.txt_data_inicio.value: dt_ini = datetime.strptime(self.txt_data_inicio.value, "%d/%m/%Y")
                if self.txt_data_fim.value: dt_fim = datetime.strptime(self.txt_data_fim.value, "%d/%m/%Y")
            except: pass

        lista = []
        for t in self.controller.transmissoes:
            dt_t = parse_date(t.data)
            if not dt_t: continue
            dt_norm = normalize_date(t.data)
            data_br = format_date_br(t.data)
            
            # 1. Filtro de Período
            if periodo == "A partir de hoje":
                if dt_norm < hoje_str: continue
            elif periodo == "Dia":
                if dt_norm != hoje_str: continue
            elif periodo == "Semana":
                inicio_semana = hoje - timedelta(days=hoje.weekday())
                fim_semana = inicio_semana + timedelta(days=6)
                if not (inicio_semana.date() <= dt_t.date() <= fim_semana.date()): continue
            elif periodo == "Mês":
                if dt_t.month != mes_f or dt_t.year != ano_f: continue
            elif periodo == "Ano":
                if dt_t.year != ano_f: continue
            elif periodo == "Personalizado":
                if dt_ini and dt_t < dt_ini: continue
                if dt_fim and dt_t > dt_fim: continue
            
            # 2. Outros Filtros
            if termo:
                match_evento = termo in t.evento.lower()
                match_resp = termo in t.responsavel.lower()
                match_data = termo in data_br
                if not (match_evento or match_resp or match_data): continue
                
            if self.dd_status.value != "Todos" and t.status != self.dd_status.value: continue
            if self.dd_tipo.value != "Todos" and t.tipo_transmissao != self.dd_tipo.value: continue
            if self.dd_modalidade.value != "Todos" and t.modalidade != self.dd_modalidade.value: continue
            
            # Oculta finalizados e cancelados apenas no estado inicial do calendário
            if is_default_view and t.status and (t.status.startswith("Finalizado") or t.status == "Cancelado"):
                continue

            lista.append(t)
        
        if self.dd_ordem.value == "Data do Evento": lista.sort(key=lambda x: (normalize_date(x.data), x.horario_inicio))
        else: lista.sort(key=lambda x: x.evento.lower())
        return lista

    def init_ui(self):
        lista = self.get_lista_filtrada(); tp = (len(lista) + self.itens_por_página - 1) // self.itens_por_página
        inicio = self.pagina_atual * self.itens_por_página; itens = lista[inicio:inicio+self.itens_por_página]
        self.txt_count.value = f"{len(lista)} eventos encontrados"; self.popular_grid(itens, tp)

    def popular_grid(self, itens, tp):
        if self.view_mode == "list":
            container_itens = ft.Column(spacing=10, expand=True)
            for t in itens: container_itens.controls.append(self.create_event_row(t))
        elif self.view_mode == "grid":
            container_itens = ft.ResponsiveRow(spacing=20, run_spacing=20)
            for t in itens: container_itens.controls.append(self.create_event_card(t))
        else: # modo calendario
            container_itens = self.create_calendar_view()
            
        if self.view_mode != "calendario":
            pag = ft.Row([ft.IconButton(ft.Icons.ARROW_BACK_IOS_NEW, on_click=lambda _: self.mudar_pagina(-1), disabled=self.pagina_atual == 0), ft.Text(f"Página {self.pagina_atual + 1} de {max(1, tp)}", weight=ft.FontWeight.BOLD), ft.IconButton(ft.Icons.ARROW_FORWARD_IOS, on_click=lambda _: self.mudar_pagina(1), disabled=self.pagina_atual >= tp - 1)], alignment=ft.MainAxisAlignment.CENTER)
        else:
            pag = ft.Container()
        
        grid_controls = []
        
        # Barra de "Selecionar Todos" (visível apenas no modo de seleção)
        if self.selection_mode and self.view_mode != "calendario":
            todos_ids = set(t.id for t in self.get_lista_filtrada())
            todos_selecionados = len(todos_ids) > 0 and todos_ids.issubset(self.selected_ids)
            alguns_selecionados = len(self.selected_ids) > 0 and not todos_selecionados
            
            cb_todos = ft.Checkbox(
                value=todos_selecionados,
                tristate=True if alguns_selecionados else False,
                fill_color=ft.Colors.CYAN_400,
                on_change=lambda e: self.toggle_select_all(e.control.value),
            )
            if alguns_selecionados:
                cb_todos.value = None  # Estado intermediário (tristate)
            
            selection_bar = ft.Container(
                content=ft.Row([
                    cb_todos,
                    ft.Text("Selecionar Todos", size=14, weight=ft.FontWeight.W_500, color=ft.Colors.WHITE),
                    ft.Container(expand=True),
                    ft.Container(
                        content=ft.Text(
                            f"{len(self.selected_ids)} de {len(todos_ids)} selecionados",
                            size=12, color=ft.Colors.WHITE70, weight=ft.FontWeight.W_500
                        ),
                        padding=ft.padding.symmetric(horizontal=12, vertical=4),
                        border_radius=12,
                        bgcolor=ft.Colors.with_opacity(0.15, ft.Colors.CYAN),
                    ),
                ], spacing=10, vertical_alignment=ft.CrossAxisAlignment.CENTER),
                padding=ft.padding.symmetric(horizontal=15, vertical=8),
                border_radius=10,
                border=ft.border.all(1, ft.Colors.CYAN_900),
                bgcolor=ft.Colors.with_opacity(0.08, ft.Colors.CYAN),
                margin=ft.margin.only(bottom=5),
            )
            grid_controls.append(selection_bar)
        
        grid_controls.extend([container_itens, ft.Container(height=20), pag])
        self.grid_container.controls = grid_controls
        try: self.update()
        except: pass

    def toggle_select_all(self, checked):
        """Seleciona ou desmarca todos os itens da lista filtrada."""
        todos_ids = set(t.id for t in self.get_lista_filtrada())
        if checked:
            self.selected_ids.update(todos_ids)
        else:
            self.selected_ids -= todos_ids
        self.btn_exportar_links.text = f"Gerar ({len(self.selected_ids)})"
        self.init_ui()

    def create_event_row(self, t):
        s_info = get_status_info(t.status)
        data_partes = formatar_data_semana(t.data).split(' - ')
        data_num = data_partes[0]
        dia_semana = data_partes[1] if len(data_partes) > 1 else ""
        
        return ft.GestureDetector(
            content=ft.Container(
                content=ft.Row([
                    # Checkbox de Seleção
                    ft.Checkbox(
                        value=t.id in self.selected_ids,
                        on_change=lambda e, tid=t.id: self.toggle_selection(tid, e.control.value),
                        visible=self.selection_mode,
                        fill_color=ft.Colors.CYAN_400,
                    ) if self.selection_mode else ft.Container(),
                    # Bloco de Data
                    ft.Container(
                        content=ft.Column([
                            ft.Text(data_num, weight=ft.FontWeight.BOLD, size=16, color=ft.Colors.BLACK),
                            ft.Text(dia_semana.upper(), size=9, color=ft.Colors.BLACK54, weight=ft.FontWeight.BOLD),
                        ], spacing=0, horizontal_alignment=ft.CrossAxisAlignment.CENTER),
                        bgcolor=s_info["color"],
                        padding=ft.padding.symmetric(horizontal=10, vertical=5),
                        border_radius=8,
                        width=100,
                        alignment=ft.Alignment(0, 0)
                    ),
                    # Horário e Modalidade
                    ft.Container(
                        content=ft.Column([
                            ft.Text(f"{t.horario_inicio} às {t.horario_fim}", weight=ft.FontWeight.BOLD, size=14, no_wrap=True),
                            ft.Text(t.modalidade or "N/A", size=11, color=ft.Colors.WHITE38, weight=ft.FontWeight.W_500),
                        ], spacing=2),
                        width=140
                    ),
                    # Evento e Tipo
                    ft.Container(
                        content=ft.Column([
                            ft.Text(t.evento, size=15, weight=ft.FontWeight.BOLD, max_lines=2, overflow=ft.TextOverflow.ELLIPSIS),
                            ft.Text(t.tipo_transmissao or "-", size=11, color=ft.Colors.CYAN_200, weight=ft.FontWeight.BOLD),
                        ], spacing=2, alignment=ft.MainAxisAlignment.CENTER),
                        width=350
                    ),
                    # Detalhes (Responsável, Local, Operador)
                    ft.Container(
                        content=ft.Column([
                            ft.Row([ft.Icon(ft.Icons.PERSON, size=14, color=ft.Colors.WHITE38), ft.Text(t.responsavel, size=12, weight=ft.FontWeight.W_500)], spacing=5),
                            ft.Row([ft.Icon(ft.Icons.LOCATION_ON, size=14, color=ft.Colors.WHITE38), ft.Text(t.local, size=11, color=ft.Colors.WHITE38, overflow=ft.TextOverflow.ELLIPSIS)], spacing=5),
                            ft.Row([ft.Icon(ft.Icons.ENGINEERING, size=14, color=ft.Colors.WHITE38), ft.Text(f"Op: {t.operador or 'Não definido'}", size=11, color=ft.Colors.WHITE38)], spacing=5),
                        ], spacing=2, alignment=ft.MainAxisAlignment.CENTER),
                        expand=True
                    ),
                    # Status
                    ft.Container(
                        content=ft.Row([
                            ft.Container(bgcolor=s_info["color"], width=8, height=8, border_radius=4),
                            ft.Text(s_info["label"], size=12, color=s_info["color"], weight=ft.FontWeight.BOLD)
                        ], spacing=8),
                        width=130
                    ),
                    # Ações
                    ft.Row([
                        ft.IconButton(ft.Icons.INFO_OUTLINE, icon_size=18, icon_color=ft.Colors.CYAN_200, tooltip="Detalhes", on_click=lambda _: mostrar_detalhes_transmissao(self.page, t, self.controller, self.atualizar)),
                        ft.IconButton(ft.Icons.EDIT_NOTE, icon_size=20, icon_color=ft.Colors.BLUE_200, tooltip="Editar", on_click=lambda _: self.on_edit(t)),
                        ft.IconButton(ft.Icons.DELETE_OUTLINE, icon_size=18, icon_color=ft.Colors.RED_300, tooltip="Excluir", on_click=lambda _: self.confirmar_exclusao(t))
                    ], spacing=0)
                ], spacing=15),
                padding=ft.padding.symmetric(horizontal=15, vertical=10),
                border_radius=12,
                border=ft.border.all(1, ft.Colors.WHITE10),
                bgcolor=ft.Colors.WHITE10,
            ),
            on_tap=lambda _: self.row_toggle_selection(t.id) if self.selection_mode else mostrar_detalhes_transmissao(self.page, t, self.controller, self.atualizar),
            on_secondary_tap=lambda e: abrir_menu_contexto(self.page, t, self.controller, self.on_edit, self.confirmar_exclusao, self.atualizar),
        )

    def create_event_card(self, t):
        s_info = get_status_info(t.status)
        return ft.Container(
            col={"sm": 12, "md": 6, "xl": 4}, 
            content=ft.GestureDetector(
                content=ft.Container(
                    content=ft.Column([
                        ft.Row([
                            ft.Checkbox(
                                value=t.id in self.selected_ids,
                                on_change=lambda e, tid=t.id: self.toggle_selection(tid, e.control.value),
                                visible=self.selection_mode,
                                fill_color=ft.Colors.CYAN_400,
                            ) if self.selection_mode else ft.Container(),
                            ft.Container(
                                content=ft.Text(format_date_br(t.data), size=12, weight=ft.FontWeight.BOLD, color=ft.Colors.BLACK), 
                                bgcolor=s_info["color"], 
                                padding=ft.padding.symmetric(horizontal=10, vertical=5), 
                                border_radius=5
                            ), 
                            ft.Container(expand=True), 
                            ft.Text(f"{t.horario_inicio} - {t.horario_fim}", size=12, weight=ft.FontWeight.BOLD)
                        ], alignment=ft.MainAxisAlignment.SPACE_BETWEEN), 
                        
                        ft.Text(t.evento, size=18, weight=ft.FontWeight.BOLD, max_lines=2, overflow=ft.TextOverflow.ELLIPSIS, height=50), 
                        
                        ft.Row([
                            ft.Container(content=ft.Text(t.tipo_transmissao or "-", size=10, weight=ft.FontWeight.BOLD, color=ft.Colors.CYAN_200), border=ft.border.all(1, ft.Colors.CYAN_900), padding=ft.padding.symmetric(horizontal=8, vertical=4), border_radius=5),
                            ft.Container(content=ft.Text(t.modalidade or "-", size=10, weight=ft.FontWeight.BOLD, color=ft.Colors.WHITE38), border=ft.border.all(1, ft.Colors.WHITE10), padding=ft.padding.symmetric(horizontal=8, vertical=4), border_radius=5),
                        ], spacing=10),

                        ft.Column([
                            ft.Row([ft.Icon(ft.Icons.PERSON, size=16, color=ft.Colors.WHITE38), ft.Text(t.responsavel, size=13, color=ft.Colors.WHITE70, expand=True)], spacing=10), 
                            ft.Row([ft.Icon(ft.Icons.LOCATION_ON, size=16, color=ft.Colors.WHITE38), ft.Text(t.local, size=13, color=ft.Colors.WHITE70, expand=True)], spacing=10),
                            ft.Row([ft.Icon(ft.Icons.ENGINEERING, size=16, color=ft.Colors.WHITE38), ft.Text(t.operador or "Sem operador", size=13, color=ft.Colors.WHITE38, expand=True)], spacing=10),
                        ], spacing=5), 
                        
                        ft.Divider(height=10, color=ft.Colors.WHITE10), 
                        
                        ft.Row([
                            ft.Container(
                                content=ft.Row([
                                    ft.Container(bgcolor=s_info["color"], width=8, height=8, border_radius=4), 
                                    ft.Text(s_info["label"], size=11, color=s_info["color"], weight=ft.FontWeight.BOLD)
                                ], spacing=5), 
                                expand=True
                            ), 
                            ft.IconButton(ft.Icons.INFO_OUTLINE, icon_size=18, icon_color=ft.Colors.CYAN_200, on_click=lambda _: mostrar_detalhes_transmissao(self.page, t, self.controller, self.atualizar)), 
                            ft.IconButton(ft.Icons.EDIT_NOTE, icon_size=18, icon_color=ft.Colors.BLUE_200, on_click=lambda _: self.on_edit(t))
                        ], alignment=ft.MainAxisAlignment.SPACE_BETWEEN)
                    ], spacing=12), 
                    padding=20, 
                    border_radius=15, 
                    bgcolor=ft.Colors.WHITE10, 
                    border=ft.border.all(1, ft.Colors.with_opacity(0.1, ft.Colors.WHITE))
                ), 
                on_tap=lambda _: self.row_toggle_selection(t.id) if self.selection_mode else mostrar_detalhes_transmissao(self.page, t, self.controller, self.atualizar), 
                on_secondary_tap=lambda e: abrir_menu_contexto(self.page, t, self.controller, self.on_edit, self.confirmar_exclusao, self.atualizar)
            )
        )

    def abrir_pasta(self, caminho):
        try:
            pasta = os.path.dirname(caminho)
            if os.name == 'nt': os.startfile(pasta)
            else: import subprocess; subprocess.run(['open' if os.uname().sysname == 'Darwin' else 'xdg-open', pasta])
        except: pass

    def toggle_selection(self, tid, checked):
        if checked: self.selected_ids.add(tid)
        else: self.selected_ids.discard(tid)
        self.btn_exportar_links.text = f"Gerar ({len(self.selected_ids)})"
        self.update()

    def row_toggle_selection(self, tid):
        if tid in self.selected_ids: self.selected_ids.discard(tid)
        else: self.selected_ids.add(tid)
        self.btn_exportar_links.text = f"Gerar ({len(self.selected_ids)})"
        self.init_ui()

    def cancelar_selecao_links(self, e):
        self.selection_mode = False
        self.selected_ids.clear()
        self.btn_exportar_links.text = "Links"
        self.btn_exportar_links.style.bgcolor = ft.Colors.BLUE_700
        self.btn_cancelar_links.visible = False
        self.init_ui()

    def exportar_excel(self, e):
        r = tk.Tk(); r.withdraw(); r.attributes("-topmost", True); r.focus_force(); path = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel", "*.xlsx")], initialfile="calendario.xlsx"); r.destroy()
        if not path: return
        lista = self.get_lista_filtrada()
        async def _ex():
            try:
                def _build():
                    wb = Workbook(); ws = wb.active; ws.title = "Calendário"
                    
                    # Cabeçalhos
                    headers = ["Data", "Início", "Fim", "Evento", "Responsável", "Local", "Tipo", "Modalidade", "Status", "Público", "Duração", "Operador", "Links", "Observações"]
                    
                    # Estilo para o cabeçalho
                    header_font = Font(bold=True, color="FFFFFF")
                    header_fill = PatternFill(start_color="333333", end_color="333333", fill_type="solid")
                    
                    for col_idx, header in enumerate(headers, 1):
                        cell = ws.cell(row=1, column=col_idx, value=header)
                        cell.font = header_font
                        cell.fill = header_fill
                        cell.alignment = Alignment(horizontal="center", vertical="center")

                    # Mapeamento de cores Flet para Hex (OpenPyXL)
                    # Usamos tons um pouco mais pastéis para o Excel
                    status_colors_hex = {
                        ft.Colors.RED_400: "FFCDD2",
                        ft.Colors.AMBER_400: "FFF9C4",
                        ft.Colors.TEAL_400: "B2DFDB",
                        ft.Colors.GREEN_400: "C8E6C9",
                        ft.Colors.PURPLE_400: "E1BEE7",
                        ft.Colors.BROWN: "D7CCC8",
                    }

                    for r_idx, t in enumerate(lista, 2):
                        # Agregando links
                        links = []
                        if t.link_youtube: links.append(f"YouTube: {t.link_youtube}")
                        if t.link_stream: links.append(f"{t.tipo_transmissao or 'Stream'}: {t.link_stream}")
                        if hasattr(t, 'links_adicionais') and t.links_adicionais:
                            for l in t.links_adicionais: links.append(f"{l.get('label', 'Link')}: {l.get('url', '')}")
                        
                        data_row = [
                            format_date_br(t.data),
                            t.horario_inicio,
                            t.horario_fim,
                            t.evento,
                            t.responsavel,
                            t.local,
                            t.tipo_transmissao,
                            t.modalidade,
                            get_status_info(t.status)["label"],
                            t.publico,
                            t.tempo_total,
                            t.operador,
                            "\n".join(links),
                            t.observacoes
                        ]
                        
                        for c_idx, value in enumerate(data_row, 1):
                            cell = ws.cell(row=r_idx, column=c_idx, value=value)
                            
                            # Formatação da coluna de Status (Índice 9)
                            if c_idx == 9:
                                s_info = get_status_info(t.status)
                                hex_color = status_colors_hex.get(s_info["color"], "FFFFFF")
                                cell.fill = PatternFill(start_color=hex_color, end_color=hex_color, fill_type="solid")
                                cell.alignment = Alignment(horizontal="center", vertical="center")
                                cell.font = Font(bold=True)
                            else:
                                cell.alignment = Alignment(vertical="center", wrap_text=(c_idx in [4, 13, 14]))

                    # Ajuste automático de largura
                    for col in ws.columns:
                        max_length = 0; column = col[0].column_letter
                        for cell in col:
                            try:
                                if cell.value:
                                    # Para células com quebra de linha, pegamos a maior linha
                                    lines = str(cell.value).split('\n')
                                    length = max(len(line) for line in lines)
                                    if length > max_length: max_length = length
                            except: pass
                        ws.column_dimensions[column].width = min(max_length + 3, 60)

                    wb.save(path)
                await asyncio.to_thread(_build)
                self.show_popup("Excel salvo com sucesso!", path)
            except Exception as ex: print(f"Erro ao exportar Excel: {ex}")
        self.page.run_task(_ex)

    def exportar_links(self, e):
        if not self.selection_mode:
            self.selection_mode = True
            self.selected_ids.clear()
            self.btn_exportar_links.text = "Gerar (0)"
            self.btn_exportar_links.style.bgcolor = ft.Colors.ORANGE_800
            self.btn_cancelar_links.visible = True
            self.init_ui()
            return

        if not self.selected_ids:
            self.page.snack_bar = ft.SnackBar(ft.Text("Selecione ao menos um item para gerar links!"), bgcolor=ft.Colors.ORANGE_800)
            self.page.snack_bar.open = True
            self.page.update()
            return

        r = tk.Tk(); r.withdraw(); r.attributes("-topmost", True); r.focus_force(); path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word", "*.docx")], initialfile="links_selecionados.docx"); r.destroy()
        if not path: return
        self.btn_exportar_links.disabled = True; self.btn_exportar_links.text = "Gerando..."; self.update()
        
        # Filtra a lista apenas com os selecionados
        lista_completa = self.get_lista_filtrada()
        lista = [t for t in lista_completa if t.id in self.selected_ids]
        
        async def _ex():
            try:
                def _process():
                    doc = Document(); title = doc.add_heading('Links das transmissões', 0); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for t in lista:
                        p_evento = doc.add_paragraph(); run_evento = p_evento.add_run(t.evento); run_evento.bold = True; run_evento.font.size = Pt(14)
                        p_data = doc.add_paragraph(); p_data.add_run(f"{format_date_br(t.data)} - {t.horario_inicio} às {t.horario_fim}")
                        if t.link_youtube: doc.add_paragraph(f"YouTube: {t.link_youtube}")
                        if t.link_stream: doc.add_paragraph(f"{t.tipo_transmissao or 'Stream'}: {t.link_stream}")
                        if hasattr(t, 'links_adicionais') and t.links_adicionais:
                            for link_obj in t.links_adicionais:
                                label = link_obj.get("label", "Link"); url = link_obj.get("url", ""); 
                                if url: doc.add_paragraph(f"{label}: {url}")
                        doc.add_paragraph("-" * 15)
                    doc.save(path)
                await asyncio.to_thread(_process)
                self.show_popup("Relatório de Links salvo!", path)
                # Reset selection mode after success
                self.selection_mode = False
                self.selected_ids.clear()
                self.btn_cancelar_links.visible = False
            except Exception as ex: print(f"Erro docx: {ex}")
            finally: self.btn_exportar_links.disabled = False; self.btn_exportar_links.text = "Links"; self.btn_exportar_links.style.bgcolor = ft.Colors.BLUE_700; self.page.update(); self.init_ui()
        self.page.run_task(_ex)

    def show_popup(self, text, path):
        def on_open_click(e):
            self.abrir_pasta(path)
            dlg.open = False
            self.page.update()
        dlg = ft.AlertDialog(
            modal=True,
            title=ft.Text("Sucesso!"),
            content=ft.Text(f"{text}\nDeseja abrir a pasta onde o arquivo foi salvo?"),
            actions=[
                ft.TextButton("Não", on_click=lambda _: self.fechar_dlg(dlg)),
                ft.ElevatedButton("Abrir Pasta", icon=ft.Icons.FOLDER_OPEN, bgcolor=ft.Colors.GREEN_700, on_click=on_open_click)
            ],
            actions_alignment=ft.MainAxisAlignment.END
        )
        self.page.overlay.append(dlg)
        dlg.open = True
        self.page.update()

    def confirmar_exclusao(self, t):
        dlg = ft.AlertDialog(title=ft.Text("Confirmar Exclusão"), content=ft.Text(f"Deseja excluir '{t.evento}'?"), actions=[ft.TextButton("Cancelar", on_click=lambda _: self.fechar_dlg(dlg)), ft.ElevatedButton("Excluir", bgcolor=ft.Colors.RED_700, on_click=lambda _: self.deletar_e_fechar(t, dlg))])
        self.page.overlay.append(dlg); dlg.open = True; self.page.update()
    def fechar_dlg(self, dlg): dlg.open = False; self.page.update()
    def deletar_e_fechar(self, t, dlg): self.controller.deletar(t.id); dlg.open = False; self.reset_and_update()
    def reset_and_update(self): 
        self.pagina_atual = 0; 
        self.refresh_dropdowns()
        self.init_ui()
        self.update_year_menu()
        self.atualizar_meses_disponiveis_bar()

    def atualizar(self): 
        self.refresh_dropdowns()
        self.init_ui()
        self.update_year_menu()
        self.atualizar_meses_disponiveis_bar()

    def refresh_dropdowns(self):
        # Atualiza as opções dos dropdowns de Tipo e Modalidade
        tipos = sorted(list(set(t.tipo_transmissao for t in self.controller.transmissoes if t.tipo_transmissao)))
        self.dd_tipo.options = [ft.dropdown.Option("Todos")] + [ft.dropdown.Option(o) for o in tipos]
        
        modalidades = sorted(list(set(t.modalidade for t in self.controller.transmissoes if t.modalidade)))
        self.dd_modalidade.options = [ft.dropdown.Option("Todos")] + [ft.dropdown.Option(o) for o in modalidades]
        
        # Manter o valor atual se ele ainda existir, senão volta para "Todos"
        if self.dd_tipo.value not in [o.key for o in self.dd_tipo.options if o.key] and self.dd_tipo.value not in [o.text for o in self.dd_tipo.options]:
            self.dd_tipo.value = "Todos"
        if self.dd_modalidade.value not in [o.key for o in self.dd_modalidade.options if o.key] and self.dd_modalidade.value not in [o.text for o in self.dd_modalidade.options]:
            self.dd_modalidade.value = "Todos"

    def set_filtro_mes_bar(self, mes_id, mes_nome):
        self.txt_mes_selecionado.value = mes_nome
        if mes_id == 0:
            self.dd_periodo.value = "Ano"
        else:
            self.dd_filtro_mes.value = str(mes_id)
            self.dd_periodo.value = "Mês"
        self.reset_and_update()

    def set_filtro_ano_bar(self, ano):
        self.txt_ano_selecionado.value = str(ano)
        self.dd_filtro_ano.value = str(ano)
        self.dd_periodo.value = "Mês"
        self.reset_and_update()

    def update_year_menu(self):
        anos_disp = sorted(list(set(parse_date(t.data).year for t in self.controller.transmissoes if parse_date(t.data))), reverse=True)
        if not anos_disp: anos_disp = [datetime.now().year]
        self.menu_ano.items = [ft.PopupMenuItem(content=ft.Text(str(a)), on_click=lambda e, a=a: self.set_filtro_ano_bar(a)) for a in anos_disp]
        if self.txt_ano_selecionado.value not in [str(a) for a in anos_disp]:
             self.txt_ano_selecionado.value = str(anos_disp[0])
             self.dd_filtro_ano.value = str(anos_disp[0])

    def atualizar_meses_disponiveis_bar(self):
        try:
            ano = int(self.txt_ano_selecionado.value)
            meses_com_dados = set()
            for t in self.controller.transmissoes:
                dt = parse_date(t.data)
                if dt and dt.year == ano:
                    meses_com_dados.add(dt.month)
            
            self.row_meses_disponiveis.controls = []
            for m_idx in range(1, 13):
                if m_idx in meses_com_dados:
                    nome_abrev = self.meses_nomes[m_idx-1][:3].upper()
                    is_selected = (m_idx == int(self.dd_filtro_mes.value)) and (self.dd_periodo.value == "Mês")
                    self.row_meses_disponiveis.controls.append(
                        ft.TextButton(
                            content=ft.Text(
                                nome_abrev, 
                                color=ft.Colors.WHITE if is_selected else ft.Colors.WHITE70, 
                                size=13, 
                                weight=ft.FontWeight.BOLD if is_selected else ft.FontWeight.W_500
                            ),
                            on_click=lambda e, m=m_idx: self.set_filtro_mes_bar(m, self.meses_nomes[m-1]),
                            style=ft.ButtonStyle(padding=ft.padding.all(0))
                        )
                    )
            self.update()
        except: pass

    def mudar_pagina(self, delta): self.pagina_atual += delta; self.init_ui()

    def create_calendar_view(self):
        mes = int(self.dd_filtro_mes.value) if self.dd_filtro_mes.value else datetime.now().month
        ano = int(self.dd_filtro_ano.value) if self.dd_filtro_ano.value else datetime.now().year
        termo = self.txt_busca.value.lower().strip() if self.txt_busca.value else ""
        
        # Filtros ativos
        filtro_status = self.dd_status.value
        filtro_tipo = self.dd_tipo.value
        filtro_modalidade = self.dd_modalidade.value
        
        # Obtém eventos do mês, aplicando filtros
        all_events = self.controller.transmissoes
        events_by_day = {}
        for t in all_events:
            dt = parse_date(t.data)
            if not dt or dt.month != mes or dt.year != ano:
                continue
            # Aplica filtros
            if filtro_status != "Todos" and t.status != filtro_status:
                continue
            if filtro_tipo != "Todos" and t.tipo_transmissao != filtro_tipo:
                continue
            if filtro_modalidade != "Todos" and t.modalidade != filtro_modalidade:
                continue
            # Aplica busca
            if termo:
                if not (termo in t.evento.lower() or termo in t.responsavel.lower() or termo in (t.local or "").lower()):
                    continue
            d = dt.day
            if d not in events_by_day:
                events_by_day[d] = []
            events_by_day[d].append(t)
        
        cal = calendar.Calendar(firstweekday=6)
        month_days = cal.monthdayscalendar(ano, mes)
        
        headers = ["DOM", "SEG", "TER", "QUA", "QUI", "SEX", "SÁB"]
        header_row = ft.Row(
            [ft.Container(
                content=ft.Text(h, weight=ft.FontWeight.BOLD, size=12, color=ft.Colors.WHITE38),
                expand=1,
                alignment=ft.Alignment(0, 0),
                padding=ft.padding.only(bottom=10)
            ) for h in headers],
            spacing=10
        )
        
        calendar_rows = [header_row]
        for week in month_days:
            week_row = ft.Row(spacing=10, expand=False)
            for day in week:
                if day == 0:
                    week_row.controls.append(ft.Container(expand=1, height=120))
                else:
                    events = events_by_day.get(day, [])
                    week_row.controls.append(self.create_day_cell(day, events, mes, ano, termo))
            calendar_rows.append(week_row)
            
        return ft.Container(
            content=ft.Column(calendar_rows, spacing=8, expand=True),
            padding=ft.padding.only(bottom=30)
        )

    def create_day_cell(self, day, events, mes, ano, termo=""):
        hoje = datetime.now()
        is_hoje = hoje.day == day and hoje.month == mes and hoje.year == ano
        
        bg_color = ft.Colors.with_opacity(0.06, ft.Colors.WHITE)
        border_color = ft.Colors.with_opacity(0.08, ft.Colors.WHITE)
        border_width = 1
        
        if is_hoje:
            border_color = ft.Colors.CYAN_400
            border_width = 2
            bg_color = ft.Colors.with_opacity(0.15, ft.Colors.CYAN_900)

        if events:
            active_events = [e for e in events if e.status and not (e.status.startswith("Finalizado") or e.status == "Cancelado")]
            target_event = active_events[0] if active_events else events[0]
            status_color = get_status_info(target_event.status)["color"]
            bg_color = ft.Colors.with_opacity(0.2, status_color)
            if not is_hoje:
                border_color = ft.Colors.with_opacity(0.5, status_color)

        # Resumo das transmissões dentro da célula
        event_summaries = ft.Column(spacing=3, tight=True)
        for t in events[:3]:
            s_info = get_status_info(t.status)
            # Destaque de busca: texto amarelo se contiver o termo
            has_match = termo and (termo in t.evento.lower() or termo in t.responsavel.lower())
            txt_color = ft.Colors.YELLOW_300 if has_match else ft.Colors.WHITE
            txt_weight = ft.FontWeight.BOLD if has_match else ft.FontWeight.W_500
            
            event_summaries.controls.append(
                ft.Container(
                    content=ft.Text(
                        f"{t.horario_inicio} {t.evento}",
                        size=11, weight=txt_weight,
                        color=txt_color, max_lines=1,
                        overflow=ft.TextOverflow.ELLIPSIS
                    ),
                    bgcolor=ft.Colors.with_opacity(0.5, s_info["color"]),
                    padding=ft.padding.symmetric(horizontal=5, vertical=2),
                    border_radius=4,
                    border=ft.border.all(1, ft.Colors.YELLOW_400) if has_match else None,
                    on_click=lambda _, item=t: mostrar_detalhes_transmissao(self.page, item, self.controller, self.atualizar, on_edit=self.on_edit),
                )
            )

        if len(events) > 3:
            event_summaries.controls.append(
                ft.Text(f"+ {len(events)-3} mais", size=10, italic=True, color=ft.Colors.WHITE54)
            )

        # Cabeçalho do dia
        day_header = ft.Row([
            ft.Text(str(day), weight=ft.FontWeight.BOLD, size=15,
                    color=ft.Colors.WHITE if events or is_hoje else ft.Colors.WHITE38),
            ft.Container(
                content=ft.Text(str(len(events)), size=9, color=ft.Colors.WHITE, weight=ft.FontWeight.BOLD),
                bgcolor=ft.Colors.with_opacity(0.3, ft.Colors.WHITE),
                border_radius=8, width=18, height=18,
                alignment=ft.Alignment(0, 0),
            ) if events else ft.Container(),
        ], alignment=ft.MainAxisAlignment.SPACE_BETWEEN)

        return ft.Container(
            content=ft.Column([day_header, event_summaries], spacing=3),
            padding=ft.padding.only(left=6, right=6, top=5, bottom=4),
            bgcolor=bg_color,
            expand=1,
            height=120,
            border_radius=10,
            border=ft.border.all(border_width, border_color),
            # Clique esquerdo na área vazia → menu de contexto
            on_click=lambda _: self.abrir_menu_dia(day, events, mes, ano),
            # Clique longo (alternativa ao botão direito) → abre lista de transmissões
            on_long_press=lambda _: self.show_day_details(day, events) if events else None,
        )

    def show_day_details(self, day, events):
        if events:
            self.mostrar_lista_dia(day, events, int(self.dd_filtro_mes.value), int(self.dd_filtro_ano.value))

    def mostrar_lista_dia(self, day, events, mes, ano):
        list_items = []
        for t in events:
            s_info = get_status_info(t.status)
            list_items.append(
                ft.Container(
                    content=ft.ListTile(
                        leading=ft.Container(width=12, height=12, border_radius=6, bgcolor=s_info["color"]),
                        title=ft.Text(f"{t.horario_inicio} - {t.evento}", weight=ft.FontWeight.BOLD, size=14),
                        subtitle=ft.Text(f"👤 {t.responsavel}  |  📍 {t.local or 'N/A'}", size=12, color=ft.Colors.WHITE38),
                        on_click=lambda _, item=t: self.abrir_detalhe_e_fechar(item, dlg),
                    ),
                    bgcolor=ft.Colors.with_opacity(0.05, ft.Colors.WHITE),
                    border_radius=10,
                    margin=ft.margin.only(bottom=5)
                )
            )
        
        def fechar(e):
            dlg.open = False
            self.page.update()

        dlg = ft.AlertDialog(
            title=ft.Row([
                ft.Icon(ft.Icons.LIST_ALT, color=ft.Colors.CYAN_400), 
                ft.Text(f"Transmissões de {day:02d}/{mes:02d}/{ano}", size=18, weight=ft.FontWeight.BOLD)
            ], spacing=10),
            content=ft.Container(
                content=ft.Column(list_items, scroll=ft.ScrollMode.AUTO, tight=True, spacing=5),
                width=550
            ),
            actions=[ft.TextButton("Fechar", on_click=fechar)],
            shape=ft.RoundedRectangleBorder(radius=15),
            bgcolor=ft.Colors.GREY_900,
        )
        self.page.overlay.append(dlg)
        dlg.open = True
        self.page.update()

    def abrir_detalhe_e_fechar(self, item, dlg):
        dlg.open = False
        self.page.update()
        mostrar_detalhes_transmissao(self.page, item, self.controller, self.atualizar)

    def abrir_menu_dia(self, day, events, mes, ano):
        from datetime import date
        dt_selecionada = date(ano, mes, day)
        
        def acao_ver(e):
            menu_dlg.open = False
            self.page.update()
            if len(events) > 1:
                self.mostrar_lista_dia(day, events, mes, ano)
            elif len(events) == 1:
                mostrar_detalhes_transmissao(self.page, events[0], self.controller, self.atualizar, on_edit=self.on_edit)

        def acao_adicionar(e):
            menu_dlg.open = False
            self.page.update()
            self.on_edit(None, default_date=dt_selecionada)

        def fechar_menu(e):
            menu_dlg.open = False
            self.page.update()

        menu_dlg = ft.AlertDialog(
            title=ft.Row([ft.Icon(ft.Icons.CALENDAR_TODAY, color=ft.Colors.CYAN_400), ft.Text(f"Dia {day:02d}/{mes:02d}")]),
            content=ft.Column([
                ft.TextButton(
                    content=ft.Row([ft.Icon(ft.Icons.LIST, color=ft.Colors.BLUE_200), ft.Text("Ver Transmissões", color=ft.Colors.WHITE, size=14)], spacing=10),
                    on_click=acao_ver,
                    disabled=not events
                ),
                ft.TextButton(
                    content=ft.Row([ft.Icon(ft.Icons.ADD_BOX, color=ft.Colors.GREEN_400), ft.Text("Adicionar Evento", color=ft.Colors.WHITE, size=14)], spacing=10),
                    on_click=acao_adicionar
                ),
            ], tight=True, spacing=10),
            actions=[ft.TextButton("Fechar", on_click=fechar_menu)],
            shape=ft.RoundedRectangleBorder(radius=15),
            bgcolor=ft.Colors.GREY_900,
        )
        self.page.overlay.append(menu_dlg)
        menu_dlg.open = True
        self.page.update()
